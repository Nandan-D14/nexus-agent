# Copyright (c) 2026 nandan-d14. All rights reserved.
# Proprietary and non-commercial use only.

"""Sandbox-side renderer for themed Word documents.

The previous generator walked Markdown line by line and dropped anything it did
not recognise, so tables arrived in Word as literal ``| a | b |`` text. This one
parses Markdown into blocks and maps each block onto real Word constructs:
styled headings, native tables with repeating header rows, shaded code blocks,
bordered callouts, hyperlinks, images, a cover page, a TOC field, and a
header/footer with page numbers.
"""

from __future__ import annotations

DOCX_RENDERER = r'''
import json
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PAYLOAD = json.loads(open(sys.argv[1], encoding="utf-8").read())
THEME = PAYLOAD["theme"]
FONTS = THEME.get("fonts") or {}
HEAD_FONT = FONTS.get("heading") or "Segoe UI"
BODY_FONT = FONTS.get("body") or "Segoe UI"
MONO_FONT = FONTS.get("mono") or "Consolas"

TITLE = str(PAYLOAD.get("title") or "Document")
SUBTITLE = str(PAYLOAD.get("subtitle") or "")
AUTHOR = str(PAYLOAD.get("author") or "")
DATE = str(PAYLOAD.get("date") or "")
OUT_PATH = PAYLOAD["out_path"]
WORKSPACE = str(PAYLOAD.get("workspace") or "")
WANT_COVER = bool(PAYLOAD.get("cover", True))
WANT_TOC = bool(PAYLOAD.get("toc", False))

MARKDOWN = str(PAYLOAD.get("markdown") or "")


def rgb(key, fallback="000000"):
    text = str(THEME.get(key) or fallback).lstrip("#").upper()
    if len(text) != 6:
        text = fallback
    return RGBColor(int(text[0:2], 16), int(text[2:4], 16), int(text[4:6], 16))


def hexval(key, fallback="000000"):
    text = str(THEME.get(key) or fallback).lstrip("#").upper()
    return text if len(text) == 6 else fallback


# ------------------------------------------------------------ xml utilities

def shade(element, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def paragraph_shading(paragraph, fill):
    shade(paragraph._p.get_or_add_pPr(), fill)


def paragraph_border(paragraph, edges):
    # edges: dict of side -> (size_eighths_pt, color, space_pt)
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        if side not in edges:
            continue
        size, col, space = edges[side]
        node = OxmlElement("w:" + side)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), str(space))
        node.set(qn("w:color"), col)
        borders.append(node)
    pPr.append(borders)


def cell_shading(cell, fill):
    shade(cell._tc.get_or_add_tcPr(), fill)


def repeat_header_row(row):
    trPr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trPr.append(node)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    return run


def add_toc_field(paragraph):
    # dirty=true makes Word rebuild the entries the first time the doc opens.
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose Update Field to build the table of contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, placeholder, end):
        run._r.append(node)


def add_hyperlink(paragraph, url, text, size):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), hexval("accent", "4F46E5"))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    for node in (fonts, col, underline, sz):
        rPr.append(node)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(rPr)
    run.append(text_node)
    link.append(run)
    paragraph._p.append(link)


# ------------------------------------------------------------ inline markup

INLINE_RE = re.compile(
    r"(\[[^\]]+?\]\([^)\s]+?\)"
    r"|\*\*[^*]+?\*\*"
    r"|__[^_]+?__"
    r"|\*[^*]+?\*"
    r"|~~[^~]+?~~"
    r"|`[^`]+?`)"
)
LINK_RE = re.compile(r"^\[([^\]]+?)\]\(([^)\s]+?)\)$")


def style_run(run, size, bold=False, italic=False, mono=False, strike=False, col=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.strike = strike
    run.font.name = MONO_FONT if mono else BODY_FONT
    run.font.color.rgb = col if col is not None else rgb("ink", "0F172A")


def add_inline(paragraph, text, size, base_color=None, bold=False, italic=False):
    col = base_color if base_color is not None else rgb("ink", "0F172A")
    for piece in [p for p in INLINE_RE.split(str(text)) if p]:
        link = LINK_RE.match(piece)
        if link:
            add_hyperlink(paragraph, link.group(2), link.group(1), size)
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) > 4:
            style_run(paragraph.add_run(piece[2:-2]), size, bold=True, italic=italic, col=col)
        elif piece.startswith("__") and piece.endswith("__") and len(piece) > 4:
            style_run(paragraph.add_run(piece[2:-2]), size, bold=True, italic=italic, col=col)
        elif piece.startswith("~~") and piece.endswith("~~") and len(piece) > 4:
            style_run(paragraph.add_run(piece[2:-2]), size, bold=bold, italic=italic,
                      strike=True, col=col)
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            style_run(paragraph.add_run(piece[1:-1]), size, bold=bold, italic=True, col=col)
        elif piece.startswith("`") and piece.endswith("`") and len(piece) > 2:
            style_run(paragraph.add_run(piece[1:-1]), size - 0.5, mono=True,
                      col=rgb("accent", "4F46E5"))
        else:
            style_run(paragraph.add_run(piece), size, bold=bold, italic=italic, col=col)


# ------------------------------------------------------------ block parsing

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
HR_RE = re.compile(r"^([-*_])\s*(?:\1\s*){2,}$")
LIST_RE = re.compile(r"^(\s*)([-*+]|\d+[.)])\s+(.*)$")
IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)\s]+)\)\s*$")
TABLE_SEP_RE = re.compile(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?$")


def split_row(line):
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def parse_blocks(text):
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks = []
    i = 0
    total = len(lines)
    while i < total:
        raw = lines[i]
        stripped = raw.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```") or stripped.startswith("~~~"):
            fence = stripped[:3]
            lang = stripped[3:].strip()
            i += 1
            buf = []
            while i < total and not lines[i].strip().startswith(fence):
                buf.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", {"lang": lang, "lines": buf}))
            continue

        heading = HEADING_RE.match(stripped)
        if heading:
            blocks.append(("heading", {
                "level": len(heading.group(1)),
                "text": heading.group(2).strip().rstrip("#").strip(),
            }))
            i += 1
            continue

        if HR_RE.match(stripped):
            blocks.append(("hr", {}))
            i += 1
            continue

        image = IMAGE_RE.match(stripped)
        if image:
            blocks.append(("image", {"alt": image.group(1), "src": image.group(2)}))
            i += 1
            continue

        if "|" in stripped and i + 1 < total and TABLE_SEP_RE.match(lines[i + 1].strip()):
            headers = split_row(raw)
            aligns = []
            for spec in split_row(lines[i + 1]):
                left = spec.startswith(":")
                right = spec.endswith(":")
                aligns.append("center" if (left and right) else ("right" if right else "left"))
            i += 2
            rows = []
            while i < total and "|" in lines[i] and lines[i].strip():
                rows.append(split_row(lines[i]))
                i += 1
            width = len(headers)
            rows = [(row + [""] * width)[:width] for row in rows]
            aligns = (aligns + ["left"] * width)[:width]
            blocks.append(("table", {"headers": headers, "rows": rows, "aligns": aligns}))
            continue

        if stripped.startswith(">"):
            buf = []
            while i < total and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append(("quote", {"lines": [b for b in buf if b]}))
            continue

        if LIST_RE.match(raw):
            items = []
            while i < total and LIST_RE.match(lines[i]):
                match = LIST_RE.match(lines[i])
                indent = len(match.group(1).replace("\t", "  "))
                marker = match.group(2)
                items.append({
                    "level": min(indent // 2, 2),
                    "ordered": marker[0].isdigit(),
                    "text": match.group(3).strip(),
                })
                i += 1
            blocks.append(("list", {"items": items}))
            continue

        buf = [stripped]
        i += 1
        while i < total:
            nxt = lines[i]
            nxt_stripped = nxt.strip()
            if (not nxt_stripped or HEADING_RE.match(nxt_stripped) or HR_RE.match(nxt_stripped)
                    or LIST_RE.match(nxt) or nxt_stripped.startswith(">")
                    or nxt_stripped.startswith("|")
                    or nxt_stripped.startswith("```") or nxt_stripped.startswith("~~~")
                    or IMAGE_RE.match(nxt_stripped)):
                break
            buf.append(nxt_stripped)
            i += 1
        blocks.append(("paragraph", {"text": " ".join(buf)}))
    return blocks


# ------------------------------------------------------------- doc assembly

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(0.9)
# Keep the running header and page number off the cover page.
section.different_first_page_header_footer = WANT_COVER
CONTENT_WIDTH = section.page_width - section.left_margin - section.right_margin

normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = rgb("ink", "0F172A")
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.25

HEADING_SPECS = {
    1: (19, "ink", True, 18, 8),
    2: (15, "accent", True, 16, 6),
    3: (12.5, "ink", True, 12, 4),
    4: (11, "muted", True, 10, 3),
}
for level, (size, key, bold, before, after) in HEADING_SPECS.items():
    try:
        style = doc.styles["Heading " + str(level)]
    except KeyError:
        continue
    style.font.name = HEAD_FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = rgb(key, "0F172A")
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True


def resolve_asset(path):
    candidate = str(path or "").strip()
    if not candidate:
        return None
    options = [candidate]
    if not candidate.startswith("/") and WORKSPACE:
        options.append(os.path.join(WORKSPACE, candidate))
    for option in options:
        if os.path.isfile(option):
            return option
    return None


def build_cover():
    accent_bar = doc.add_paragraph()
    accent_bar.paragraph_format.space_before = Pt(120)
    accent_bar.paragraph_format.space_after = Pt(0)
    paragraph_border(accent_bar, {"bottom": (36, hexval("accent", "4F46E5"), 0)})

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run(TITLE)
    run.font.name = HEAD_FONT
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = rgb("ink", "0F172A")

    if SUBTITLE:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(4)
        run = sub.add_run(SUBTITLE)
        run.font.name = BODY_FONT
        run.font.size = Pt(13)
        run.font.color.rgb = rgb("muted", "64748B")

    meta = "   |   ".join([v for v in (AUTHOR, DATE) if v])
    if meta:
        line = doc.add_paragraph()
        line.paragraph_format.space_before = Pt(24)
        run = line.add_run(meta)
        run.font.name = BODY_FONT
        run.font.size = Pt(10)
        run.font.color.rgb = rgb("muted", "64748B")


def build_chrome():
    header = section.header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.text = ""
    run = para.add_run(TITLE)
    run.font.name = BODY_FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb("muted", "64748B")
    paragraph_border(para, {"bottom": (6, hexval("hairline", "DDE3F0"), 4)})

    footer = section.footer
    fpara = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fpara.text = ""
    fpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = fpara.add_run("Page ")
    prefix.font.name = BODY_FONT
    prefix.font.size = Pt(8.5)
    prefix.font.color.rgb = rgb("muted", "64748B")
    page_run = add_field(fpara, " PAGE ")
    page_run.font.name = BODY_FONT
    page_run.font.size = Pt(8.5)
    page_run.font.color.rgb = rgb("muted", "64748B")
    middle = fpara.add_run(" of ")
    middle.font.name = BODY_FONT
    middle.font.size = Pt(8.5)
    middle.font.color.rgb = rgb("muted", "64748B")
    total_run = add_field(fpara, " NUMPAGES ")
    total_run.font.name = BODY_FONT
    total_run.font.size = Pt(8.5)
    total_run.font.color.rgb = rgb("muted", "64748B")


def render_table(spec):
    headers = spec["headers"]
    rows = spec["rows"]
    aligns = spec["aligns"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = doc.styles["Table Grid"]
    except KeyError:
        pass
    table.autofit = True

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }

    header_cells = table.rows[0].cells
    for idx, label in enumerate(headers):
        cell = header_cells[idx]
        cell.text = ""
        cell_shading(cell, hexval("accent", "4F46E5"))
        para = cell.paragraphs[0]
        para.alignment = align_map[aligns[idx]]
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(label)
        run.font.name = BODY_FONT
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    repeat_header_row(table.rows[0])

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.text = ""
            if row_index % 2 == 1:
                cell_shading(cell, hexval("surface", "F4F6FE"))
            para = cell.paragraphs[0]
            para.alignment = align_map[aligns[idx]]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            add_inline(para, value, 9.5)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def render_code(spec):
    lines = spec["lines"] or [""]
    for index, line in enumerate(lines):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6 if index == 0 else 0)
        para.paragraph_format.space_after = Pt(6 if index == len(lines) - 1 else 0)
        para.paragraph_format.line_spacing = 1.05
        para.paragraph_format.left_indent = Inches(0.16)
        paragraph_shading(para, hexval("surface", "F4F6FE"))
        edges = {"left": (18, hexval("accent", "4F46E5"), 6)}
        if index == 0:
            edges["top"] = (4, hexval("hairline", "DDE3F0"), 2)
        if index == len(lines) - 1:
            edges["bottom"] = (4, hexval("hairline", "DDE3F0"), 2)
        paragraph_border(para, edges)
        run = para.add_run(line if line.strip() else " ")
        run.font.name = MONO_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = rgb("ink", "0F172A")


def render_quote(spec):
    lines = spec["lines"] or [""]
    for index, line in enumerate(lines):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.22)
        para.paragraph_format.space_before = Pt(8 if index == 0 else 0)
        para.paragraph_format.space_after = Pt(8 if index == len(lines) - 1 else 0)
        paragraph_shading(para, hexval("accent_soft", "E5E8FD"))
        edges = {"left": (24, hexval("accent", "4F46E5"), 8)}
        paragraph_border(para, edges)
        add_inline(para, line, 10.5, base_color=rgb("ink", "0F172A"), italic=True)


def render_list(spec):
    ordered_style = "List Number"
    bullet_style = "List Bullet"
    available = {s.name for s in doc.styles}
    for item in spec["items"]:
        level = item["level"]
        base = ordered_style if item["ordered"] else bullet_style
        name = base if level == 0 else base + " " + str(level + 1)
        if name not in available:
            name = base if base in available else None
        para = doc.add_paragraph(style=name) if name else doc.add_paragraph()
        if name is None:
            para.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
        para.paragraph_format.space_after = Pt(3)
        add_inline(para, item["text"], 10.5)


def render_image(spec):
    path = resolve_asset(spec["src"])
    if path is None:
        para = doc.add_paragraph()
        add_inline(para, "[image not found: " + str(spec["src"]) + "]", 9.5,
                   base_color=rgb("muted", "64748B"), italic=True)
        return
    try:
        doc.add_picture(path, width=CONTENT_WIDTH)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        return
    if spec["alt"]:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(spec["alt"])
        run.font.name = BODY_FONT
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = rgb("muted", "64748B")


def render_hr():
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(10)
    paragraph_border(para, {"bottom": (6, hexval("hairline", "DDE3F0"), 1)})


blocks = parse_blocks(MARKDOWN)

if WANT_COVER:
    build_cover()
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

if WANT_TOC:
    toc_heading = doc.add_paragraph()
    run = toc_heading.add_run("Contents")
    run.font.name = HEAD_FONT
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = rgb("accent", "4F46E5")
    toc_heading.paragraph_format.space_after = Pt(8)
    add_toc_field(doc.add_paragraph())
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

build_chrome()

heading_count = 0
word_count = 0
for kind, spec in blocks:
    if kind == "heading":
        heading_count += 1
        level = min(spec["level"], 4)
        para = doc.add_heading(level=level)
        add_inline(para, spec["text"], HEADING_SPECS[level][0],
                   base_color=rgb(HEADING_SPECS[level][1], "0F172A"),
                   bold=HEADING_SPECS[level][2])
        for run in para.runs:
            run.font.name = HEAD_FONT
        if level <= 2:
            paragraph_border(para, {"bottom": (4, hexval("hairline", "DDE3F0"), 4)})
        word_count += len(spec["text"].split())
    elif kind == "paragraph":
        para = doc.add_paragraph()
        add_inline(para, spec["text"], 10.5)
        word_count += len(spec["text"].split())
    elif kind == "list":
        render_list(spec)
        word_count += sum(len(item["text"].split()) for item in spec["items"])
    elif kind == "table":
        render_table(spec)
        word_count += sum(len(" ".join(row).split()) for row in spec["rows"])
    elif kind == "code":
        render_code(spec)
    elif kind == "quote":
        render_quote(spec)
        word_count += sum(len(line.split()) for line in spec["lines"])
    elif kind == "image":
        render_image(spec)
    elif kind == "hr":
        render_hr()

try:
    doc.core_properties.title = TITLE
    if AUTHOR:
        doc.core_properties.author = AUTHOR
    if SUBTITLE:
        doc.core_properties.subject = SUBTITLE
    doc.core_properties.comments = "Generated by CoComputer"
except Exception:
    pass

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)

print(json.dumps({
    "status": "success",
    "path": OUT_PATH,
    "size": os.path.getsize(OUT_PATH),
    "block_count": len(blocks),
    "heading_count": heading_count,
    "word_count": word_count,
    "table_count": sum(1 for kind, _ in blocks if kind == "table"),
    "theme": THEME.get("id") or "aurora",
}))
'''
